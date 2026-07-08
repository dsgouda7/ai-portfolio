"""DOCX extractor using python-docx."""

from __future__ import annotations

from pathlib import Path

from .base import BaseExtractor


class DocxExtractor(BaseExtractor):
    def extract(self, path: Path) -> str:
        return self._safe(self._read, path)

    def _read(self, path: Path) -> str:
        import docx  # type: ignore[import]

        doc = docx.Document(str(path))
        parts: list[str] = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)

        # Include table cells
        for table in doc.tables:
            for row in table.rows:
                row_parts = [
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                ]
                if row_parts:
                    parts.append(" | ".join(row_parts))

        return "\n\n".join(parts)
